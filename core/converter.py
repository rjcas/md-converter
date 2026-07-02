"""Motor de conversión PDF/DOCX -> Markdown.

markitdown es el motor principal. Si falla (archivo mal formado que
markitdown no tolera pero sí toleran librerías más permisivas), se
intenta un fallback específico por formato antes de rendirse.
"""
from __future__ import annotations

import errno
import logging
import os
from dataclasses import dataclass
from pathlib import Path

from markitdown import MarkItDown

from .exceptions import (
    CorruptFileError,
    EmptyContentError,
    FileLockedError,
    UnsupportedFormatError,
)

logger = logging.getLogger("md_converter.converter")

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}


@dataclass
class RawConversionResult:
    source_path: Path
    raw_markdown: str  # salida cruda de markitdown, sin optimizar


class MarkitdownConverter:
    """Wrapper delgado sobre MarkItDown con fallbacks y detección de bloqueo."""

    def __init__(self) -> None:
        # Una sola instancia reutilizable: evita reinicializar plugins en cada archivo.
        self._engine = MarkItDown()

    def convert(self, file_path: str | Path) -> RawConversionResult:
        path = Path(file_path)
        self._validate_pre_checks(path)

        try:
            result = self._engine.convert(str(path))
            text = (result.text_content or "").strip()
        except PermissionError as exc:
            raise FileLockedError(str(path), "Archivo bloqueado por otro proceso") from exc
        except Exception as primary_exc:  # noqa: BLE001 - markitdown no expone una jerarquía propia
            logger.warning("markitdown falló para %s: %s. Probando fallback.", path, primary_exc)
            text = self._fallback(path, primary_exc)

        if not text:
            raise EmptyContentError(
                str(path),
                "La conversión resultó en contenido vacío (¿PDF escaneado sin texto?)",
            )

        return RawConversionResult(source_path=path, raw_markdown=text)

    # ------------------------------------------------------------------ #
    # Validaciones y fallbacks
    # ------------------------------------------------------------------ #

    def _validate_pre_checks(self, path: Path) -> None:
        if not path.exists():
            raise CorruptFileError(str(path), "El archivo no existe")

        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            raise UnsupportedFormatError(
                str(path), f"Extensión no soportada: '{path.suffix}'"
            )

        if path.stat().st_size == 0:
            raise CorruptFileError(str(path), "El archivo está vacío (0 bytes)")

        if not os.access(path, os.R_OK):
            raise FileLockedError(str(path), "Sin permisos de lectura")

        # Detección de bloqueo real en Windows/Office: intento de apertura exclusiva.
        try:
            with open(path, "rb"):
                pass
        except OSError as exc:
            if exc.errno in (errno.EACCES, errno.EBUSY):
                raise FileLockedError(str(path), "Archivo bloqueado o en uso") from exc
            raise CorruptFileError(str(path), f"No se pudo abrir: {exc}") from exc

    def _fallback(self, path: Path, primary_exc: Exception) -> str:
        suffix = path.suffix.lower()
        try:
            if suffix == ".docx":
                return self._fallback_docx(path)
            if suffix == ".pdf":
                return self._fallback_pdf(path)
        except Exception as fallback_exc:  # noqa: BLE001
            raise CorruptFileError(
                str(path),
                f"markitdown y fallback fallaron. "
                f"Primario: {primary_exc}. Fallback: {fallback_exc}",
            ) from fallback_exc
        raise CorruptFileError(str(path), f"Sin fallback disponible: {primary_exc}")

    @staticmethod
    def _fallback_docx(path: Path) -> str:
        import docx  # python-docx, import perezoso: solo se paga el costo si se usa

        document = docx.Document(str(path))
        parts = []
        for para in document.paragraphs:
            if para.text.strip():
                style = (para.style.name or "").lower()
                if "heading 1" in style:
                    parts.append(f"# {para.text.strip()}")
                elif "heading 2" in style:
                    parts.append(f"## {para.text.strip()}")
                elif style.startswith("heading"):
                    parts.append(f"### {para.text.strip()}")
                else:
                    parts.append(para.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                parts.append(" | ".join(cells))
        return "\n\n".join(parts)

    @staticmethod
    def _fallback_pdf(path: Path) -> str:
        import pdfplumber  # import perezoso

        parts = []
        with pdfplumber.open(str(path)) as pdf:
            for page in pdf.pages:
                text = page.extract_text() or ""
                if text.strip():
                    parts.append(text.strip())
        return "\n\n".join(parts)
